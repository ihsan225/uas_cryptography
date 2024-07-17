from docx import Document
from cryptography.fernet import Fernet
import os

def generate_key():
    return Fernet.generate_key()

def save_key(key, filename):
    with open(filename, 'wb') as key_file:
        key_file.write(key)

def load_key(filename):
    with open(filename, 'rb') as key_file:
        return key_file.read()

def encrypt_file(file_path, key):
    # Load the document
    doc = Document(file_path)
    encrypted_content = ""
    
    # Read the document content
    for paragraph in doc.paragraphs:
        encrypted_content += paragraph.text + "\n"

    # Encrypt the content
    fernet = Fernet(key)
    encrypted_data = fernet.encrypt(encrypted_content.encode())

    # Save the encrypted content back to a file
    with open(file_path + ".enc", 'wb') as encrypted_file:
        encrypted_file.write(encrypted_data)
    
    os.remove(file_path)

def decrypt_file(file_path, key):
    # Load the encrypted file
    with open(file_path, 'rb') as encrypted_file:
        encrypted_data = encrypted_file.read()
    
    # Decrypt the content
    fernet = Fernet(key)
    decrypted_content = fernet.decrypt(encrypted_data).decode()

    # Create a new document and add the decrypted content
    doc = Document()
    for line in decrypted_content.split('\n'):
        doc.add_paragraph(line)

    # Save the decrypted content back to a .docx file
    new_file_path = file_path.replace('.enc', '')
    doc.save(new_file_path)
    
    os.remove(file_path)

def main():
    while True:
        print("Selamat datang di Aplikasi Enkripsi dan Dekripsi File .docx")
        print("1: Enkripsi File")
        print("2: Dekripsi File")
        print("3: Keluar")

        choice = input("Pilih opsi (1/2/3): ")

        if choice == '1':
            file_path = input("Masukkan path file .docx yang ingin dienkripsi: ")
            if not os.path.exists(file_path):
                print("File tidak ditemukan!")
                continue
            key = generate_key()
            save_key(key, 'filekey.key')
            encrypt_file(file_path, key)
            print("File berhasil dienkripsi. Kunci enkripsi disimpan di 'filekey.key'")
        elif choice == '2':
            file_path = input("Masukkan path file .docx yang ingin didekripsi (dengan ekstensi .enc): ")
            if not os.path.exists(file_path):
                print("File tidak ditemukan!")
                continue
            key = load_key('filekey.key')
            decrypt_file(file_path, key)
            print("File berhasil didekripsi.")
        elif choice == '3':
            print("Terima kasih telah menggunakan aplikasi ini!")
            break
        else:
            print("Pilihan tidak valid, silakan coba lagi.")

if __name__ == "__main__":
    main()
